"""
build_paper.py - render PHASE3_PAPER.md to a formatted .docx and .pdf for submission.

Markdown -> .docx (Times New Roman 11pt, single-spaced, 0.75" top/bottom, 0.5" side margins) via python-docx,
then -> .pdf via LibreOffice headless. Keeps PHASE3_PAPER.md as the single source of truth.

The official GIC_2026 cover page (downloaded from Aqora) must be prepended as page 1 of
the final submission - per the rules it may not be recreated/modified, so it is NOT
generated here.

Usage:  python3 build_paper.py   ->  PHASE3_PAPER.docx, PHASE3_PAPER.pdf
"""
import os
import re
import subprocess
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "PHASE3_PAPER.md")
DOCX = os.path.join(HERE, "PHASE3_PAPER.docx")


def _ensure(pkg, mod=None):
    try:
        __import__(mod or pkg)
    except ImportError:
        subprocess.run([sys.executable, "-m", "pip", "install", "--quiet", pkg], check=True)


def merge_paragraph_lines(lines):
    """Join hard-wrapped body lines into single paragraphs (markdown semantics), so
    **bold**/*italic* spans that cross a source line-wrap render correctly."""
    out, buf = [], []

    def flush():
        if buf:
            out.append(" ".join(buf)); buf.clear()
    for ln in lines:
        if ln.strip() == "" or ln.startswith(("#", "|", "- ", ">", "![")):
            flush(); out.append(ln)
        else:
            buf.append(ln.strip())
    flush()
    return out


def _add_runs(paragraph, text):
    """Add text with **bold**, *italic* and `code` inline formatting."""
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\s][^*]*\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Courier New"
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:
            paragraph.add_run(tok)


def main():
    _ensure("python-docx", "docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # base style
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.0
    st.paragraph_format.space_after = Pt(2)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.75)
        s.left_margin = s.right_margin = Inches(0.5)

    lines = merge_paragraph_lines(open(MD, encoding="utf-8").read().splitlines())
    i = 0
    while i < len(lines):
        ln = lines[i]
        if re.match(r"^> ", ln):           # drop internal authoring blockquote
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                i += 1
            continue
        if ln.strip() == "<!-- pagebreak -->":
            doc.add_page_break()
            i += 1
            continue
        if ln.startswith("| "):            # table block
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in tbl_lines]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]  # drop --- sep
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            cp = t.rows[ri].cells[ci].paragraphs[0]
                            _add_runs(cp, cell)
                            for run in cp.runs:
                                run.font.size = Pt(9); run.bold = (ri == 0) or run.bold
            continue
        m = re.match(r"^!\[(.*)\]\((.*)\)\s*$", ln)
        if m:
            cap, path = m.group(1), m.group(2)
            try:
                doc.add_picture(os.path.join(HERE, path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(cp, cap)
            for rr in cp.runs:
                rr.italic = True; rr.font.size = Pt(9)
            i += 1
            continue
        if ln.startswith("#### "):
            p = doc.add_paragraph(); r = p.add_run(ln[5:]); r.italic = True; r.bold = True
        elif ln.startswith("### "):
            doc.add_heading(ln[4:], level=3)
        elif ln.startswith("## "):
            doc.add_heading(ln[3:], level=2)
        elif ln.startswith("# "):
            doc.add_heading(ln[2:], level=1)
        elif ln.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); _add_runs(p, ln[2:])
        elif ln.strip() == "":
            pass
        else:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; _add_runs(p, ln)
        i += 1

    doc.save(DOCX)
    print(f"wrote {DOCX}")
    build_pdf(lines)


CODE_FACE = "Courier"     # switched to FreeMono (full Unicode) when freefont registers


def _inline(s):
    """markdown bold/italic/code -> reportlab markup (after escaping & < >)."""
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", s)
    s = re.sub(r"`(.+?)`", r'<font face="%s" size="10">\1</font>' % CODE_FACE, s)
    return s


def build_pdf(lines):
    """Render the same markdown to PDF via reportlab (no system deps)."""
    _ensure("reportlab")
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY

    # Full-Unicode serif (Times-like) so quantum notation (⟨Z⟩, superscripts, Greek)
    # renders instead of falling back to Latin-1 tofu boxes.
    serif, serif_b, serif_bi = "Times-Roman", "Times-Bold", "Times-BoldItalic"
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        FF = "/usr/share/fonts/truetype/freefont/"
        pdfmetrics.registerFont(TTFont("FreeSerif", FF + "FreeSerif.ttf"))
        pdfmetrics.registerFont(TTFont("FreeSerif-Bold", FF + "FreeSerifBold.ttf"))
        pdfmetrics.registerFont(TTFont("FreeSerif-Italic", FF + "FreeSerifItalic.ttf"))
        pdfmetrics.registerFont(TTFont("FreeSerif-BoldItalic", FF + "FreeSerifBoldItalic.ttf"))
        registerFontFamily("FreeSerif", normal="FreeSerif", bold="FreeSerif-Bold",
                           italic="FreeSerif-Italic", boldItalic="FreeSerif-BoldItalic")
        pdfmetrics.registerFont(TTFont("FreeMono", FF + "FreeMono.ttf"))
        global CODE_FACE
        CODE_FACE = "FreeMono"
        serif, serif_b, serif_bi = "FreeSerif", "FreeSerif-Bold", "FreeSerif-BoldItalic"
    except Exception:
        pass                               # fall back to built-in Times (Latin-1 only)

    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["Normal"], fontName=serif,
                          fontSize=11, leading=12.6, alignment=TA_JUSTIFY, spaceAfter=2)
    h1 = ParagraphStyle("h1", parent=body, fontName=serif_b, fontSize=14,
                        leading=16, alignment=0, spaceAfter=2)
    h2 = ParagraphStyle("h2", parent=body, fontName=serif_b, fontSize=12,
                        leading=14, alignment=0, spaceBefore=3, spaceAfter=2)
    h3 = ParagraphStyle("h3", parent=body, fontName=serif_b, fontSize=11,
                        leading=13, alignment=0, spaceAfter=1)
    h4 = ParagraphStyle("h4", parent=h3, fontName=serif_bi)
    note = ParagraphStyle("note", parent=body, fontSize=9, textColor=colors.grey)

    story = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("> "):
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                i += 1
            continue
        if ln.strip() == "<!-- pagebreak -->":
            from reportlab.platypus import PageBreak
            story.append(PageBreak())
            i += 1
            continue
        if ln.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in tbl]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                cellstyle = ParagraphStyle("cell", parent=body, fontSize=8.5,
                                           leading=10, alignment=1)
                data = [[Paragraph(_inline(c), cellstyle) for c in r] for r in rows]
                t = Table(data, hAlign="LEFT")
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), serif_b),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
                story.append(t); story.append(Spacer(1, 4))
            continue
        m = re.match(r"^!\[(.*)\]\((.*)\)\s*$", ln)
        if m:
            from reportlab.platypus import Image
            from PIL import Image as PILImage
            cap, path = m.group(1), m.group(2)
            ip = os.path.join(HERE, path)
            iw_px, ih_px = PILImage.open(ip).size
            iw = 5.0 * inch; ih = iw * ih_px / iw_px
            img = Image(ip, width=iw, height=ih); img.hAlign = "CENTER"
            capstyle = ParagraphStyle("cap", parent=body, fontSize=8.5, leading=10,
                                      alignment=1, textColor=colors.grey)
            story.append(img)
            story.append(Paragraph(_inline(cap), capstyle))
            story.append(Spacer(1, 4))
            i += 1
            continue
        if ln.startswith("#### "):
            story.append(Paragraph(_inline(ln[5:]), h4))
        elif ln.startswith("### "):
            story.append(Paragraph(_inline(ln[4:]), h3))
        elif ln.startswith("## "):
            story.append(Paragraph(_inline(ln[3:]), h2))
        elif ln.startswith("# "):
            story.append(Paragraph(_inline(ln[2:]), h1))
        elif ln.startswith("- "):
            story.append(Paragraph("&bull;&nbsp;" + _inline(ln[2:]), body))
        elif ln.strip() == "":
            story.append(Spacer(1, 2.5))
        else:
            story.append(Paragraph(_inline(ln), body))
        i += 1

    out = os.path.join(HERE, "PHASE3_PAPER.pdf")
    doc = SimpleDocTemplate(out, pagesize=letter, topMargin=0.75 * inch,
                            bottomMargin=0.75 * inch, leftMargin=0.5 * inch,
                            rightMargin=0.5 * inch)
    pages = []
    doc.build(story, onLaterPages=lambda c, d: pages.append(1),
              onFirstPage=lambda c, d: pages.append(1))
    print(f"wrote {out}  ({os.path.getsize(out)//1024} KB, ~{len(pages)} pages)")


if __name__ == "__main__":
    main()
